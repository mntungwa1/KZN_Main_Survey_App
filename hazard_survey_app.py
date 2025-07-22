import streamlit as st
import pandas as pd
import geopandas as gpd
import os, shutil, re, smtplib
from datetime import datetime, date
from shapely.geometry import Point
from streamlit_folium import st_folium
from email.message import EmailMessage
from dotenv import load_dotenv
import folium
from docx import Document
from fpdf import FPDF
from pathlib import Path

# --- Load environment variables ---
load_dotenv()
EMAIL_ADDRESS = os.getenv("EMAIL_ADDRESS", "dummy_email@gmail.com")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD", "dummy_password")
APP_PASSWORD = os.getenv("APP_PASSWORD", "kzn!23@")
ADMIN_EMAILS = [os.getenv("ADMIN_EMAIL", EMAIL_ADDRESS), "mhugo@srk.co.za"]

# --- Password protection ---
def password_protection():
    password = st.text_input("Enter password to access the app:", type="password")
    if password == APP_PASSWORD:
        st.session_state["authenticated"] = True
        st.success("Access granted.")
        st.rerun()
    elif password:
        st.error("Incorrect password.")

if "authenticated" not in st.session_state or not st.session_state["authenticated"]:
    st.title("KZN Hazard Risk Assessment Survey - Login")
    password_protection()
    st.stop()

# --- Email Sending ---
def send_email(subject, body, to_emails, attachments):
    try:
        msg = EmailMessage()
        msg["Subject"] = subject
        msg["From"] = EMAIL_ADDRESS
        msg["To"] = ", ".join(to_emails)
        msg.set_content(body)
        for attachment in attachments:
            with open(attachment, "rb") as f:
                file_data = f.read()
                file_name = Path(attachment).name
            msg.add_attachment(file_data, maintype="application", subtype="octet-stream", filename=file_name)
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            smtp.send_message(msg)
        st.success(f"Email sent to {to_emails}!")
    except Exception as e:
        st.error(f"Failed to send email: {e}")

# --- File paths ---
BASE_DIR = Path("C:/tmp/kzn")
TODAY_FOLDER = datetime.now().strftime("%d_%b_%Y")
SAVE_DIR = BASE_DIR / TODAY_FOLDER
SAVE_DIR.mkdir(parents=True, exist_ok=True)
EXCEL_PATH = Path("RiskAssessmentTool.xlsm")
GEOJSON_PATH = Path("KZN_wards.geojson")
LOGO_PATH = "Logo.png"
SRK_LOGO_PATH = "SRK_Logo.png"

# --- Cleanup old folders ---
def cleanup_old_folders(base_dir, days=30):
    now = datetime.now()
    pattern = re.compile(r"\d{2}_[A-Za-z]{3}_\d{4}")
    for folder in base_dir.iterdir():
        if folder.is_dir() and pattern.fullmatch(folder.name):
            try:
                if (now - datetime.strptime(folder.name, "%d_%b_%Y")).days > days:
                    shutil.rmtree(folder)
            except:
                continue
cleanup_old_folders(BASE_DIR)

# --- Load data ---
@st.cache_data(show_spinner=False)
def load_hazards():
    df = pd.read_excel(EXCEL_PATH, sheet_name="Hazard information", skiprows=1)
    hazards_list = df.iloc[:, 0].dropna().tolist()
    if not hazards_list:
        st.error("No hazards found in the 'Hazard information' sheet. Please check the Excel file.")
        st.stop()
    return hazards_list

@st.cache_data(show_spinner=False)
def load_ward_gdf():
    return gpd.read_file(GEOJSON_PATH).to_crs(epsg=4326)

hazards = load_hazards()
gdf = load_ward_gdf()

# --- Save outputs ---
def save_responses(responses, name, ward, email, date_filled):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_filename = f"{ward}_{name}_{timestamp}".replace(" ", "_")
    csv_path = SAVE_DIR / f"{base_filename}.csv"
    pdf_path = SAVE_DIR / f"{base_filename}.pdf"
    docx_path = SAVE_DIR / f"{base_filename}.docx"

    df = pd.DataFrame(responses)
    df.insert(0, "Respondent Name", name)
    df.insert(1, "Ward", ward)
    df.insert(2, "Email", email)
    df.insert(3, "Date", date_filled)

    df.to_csv(csv_path, index=False)

    doc = Document()
    doc.add_heading("KZN Hazard Risk Assessment Survey", 0)
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Ward: {ward}")
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Date: {date_filled}")
    doc.add_paragraph("---")
    for _, row in df.iterrows():
        doc.add_paragraph(f"Hazard: {row['Hazard']} | Question: {row['Question']} | Response: {row['Response']}")
    doc.save(docx_path)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(200, 10, txt="KZN Hazard Risk Assessment Survey", ln=True, align="C")
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=f"Name: {name}\nWard: {ward}\nEmail: {email}\nDate: {date_filled}\n---")
    for _, row in df.iterrows():
        pdf.multi_cell(0, 10, txt=f"Hazard: {row['Hazard']} | Question: {row['Question']} | Response: {row['Response']}")
    pdf.output(pdf_path)

    st.info(f"Files saved successfully in {SAVE_DIR} as {base_filename}.*")
    return csv_path, docx_path, pdf_path

# --- Page layout ---
st.set_page_config(page_title="KZN Hazard Risk Assessment", layout="wide")
if os.path.exists(LOGO_PATH):
    st.image(LOGO_PATH, width=240)
if os.path.exists(SRK_LOGO_PATH):
    st.image(SRK_LOGO_PATH, width=200)
st.title("KZN Hazard Risk Assessment Survey")
st.markdown("---")

# --- Interactive Map ---
st.subheader("Select a Ward from the Map")
m = folium.Map(location=[-29.5, 31.1], zoom_start=7)
folium.GeoJson(
    data=gdf.__geo_interface__,
    name="Wards",
    tooltip=folium.GeoJsonTooltip(fields=gdf.columns[:1].tolist()),
    popup=folium.GeoJsonPopup(fields=gdf.columns[:1].tolist()),
    highlight_function=lambda x: {"fillColor": "#ffaf00", "color": "black", "weight": 2},
).add_to(m)
map_data = st_folium(m, height=500)

# --- Ward selection ---
clicked_ward = None
if map_data.get("last_active_drawing"):
    props = map_data["last_active_drawing"].get("properties", {})
    clicked_ward = props.get(gdf.columns[0])
elif map_data.get("last_clicked"):
    lng = map_data["last_clicked"]["lng"]
    lat = map_data["last_clicked"]["lat"]
    pt = Point(lng, lat)
    for _, row in gdf.iterrows():
        if row.geometry.contains(pt):
            clicked_ward = row[gdf.columns[0]]
            break
if clicked_ward:
    st.session_state["selected_ward"] = clicked_ward
ward_display = st.session_state.get("selected_ward", "")
if ward_display:
    st.success(f"Selected Ward: {ward_display}")

# --- Hazard Selection ---
st.markdown("---")
st.subheader("Select Applicable Hazards")
selected = st.multiselect("Choose hazards:", hazards)
include_other = st.checkbox("Add a custom hazard (Other)")
custom = st.text_input("Specify other hazard:") if include_other else ""

if selected or custom:
    if st.button("Proceed"):
        st.session_state["show_form"] = True

questions_with_descriptions = {
    "Has this hazard occurred in the past?": [
        "0 - Has not occurred and has no chance of occurrence",
        "1 - Has not occurred but there is real potential for occurrence",
        "2 - Has occurred but only once",
        "3 - Has occurred but only a few times or rarely",
        "4 - Has occurred regularly or at least once a year",
        "5 - Occurs multiple times during a single year"
    ],
    "How frequently does it occur?": [
        "0 - Unknown / Not applicable",
        "1 - Decreasing",
        "2 - Stable",
        "3 - Marginally increasing",
        "4 - Increasing",
        "5 - Increasing rapidly"
    ],
    "What is the typical duration of the hazard?": [
        "0 - Unknown / Not applicable",
        "1 - Few Minutes",
        "2 - Few Hours",
        "3 - Few days",
        "4 - Few weeks",
        "5 - Few months"
    ],
    "What is the area of impact?": [
        "0 - None",
        "1 - Single property",
        "2 - Single Ward",
        "3 - Few wards",
        "4 - Entire municipality",
        "5 - Larger than municipality"
    ],
    "What is the impact on people?": [
        "0 - None",
        "1 - Low impact / Discomfort",
        "2 - Minimal impact / minor injuries",
        "3 - Serious injuries / health problems no fatalities",
        "4 - Fatalities / Serious health problems but confined",
        "5 - Multiple fatalities spread over wide area"
    ],
    "What is the impact on infrastructure and services?": [
        "0 - None",
        "1 - Low impact / Minor damage / Minor disruption",
        "2 - Some structural damage / Short term disruption of services",
        "3 - Medium structural damage / 1 Week disruption",
        "4 - Serious structural damage / Disruption of longer than a week",
        "5 - Total disruption of structure / Disruption of longer than a month"
    ],
    "What is the impact on the environment?": [
        "0 - Not applicable / No effects",
        "1 - Minor effects",
        "2 - Medium effects",
        "3 - Severe",
        "4 - Severe effects over wide area",
        "5 - Total Destruction"
    ],
    "What is the level of economic disruption?": [
        "0 - No disruption",
        "1 - Some disruption",
        "2 - Medium disruption",
        "3 - Severe short-term disruption",
        "4 - Severe long-term disruption",
        "5 - Total stop in activities"
    ],
    "How predictable is the hazard?": [
        "0 - Not applicable",
        "1 - Effective early warning",
        "3 - Partially predictable",
        "5 - No early warning"
    ],
    "What is the urgency or priority level?": [
        "0 - Not applicable / No effects",
        "1 - Low priority",
        "2 - Medium priority",
        "3 - Medium high priority",
        "4 - High priority",
        "5 - Very High priority"
    ]
}
capacity_questions = [
    "Sufficient staff/human resources",
    "Experience and special knowledge",
    "Equipment",
    "Funding",
    "Facilities",
    "Prevention and mitigation plans",
    "Response and recovery plans"
]
capacity_options = ["Strongly Disagree", "Disagree", "Neutral", "Agree", "Strongly Agree"]

# --- Survey Form ---
if st.session_state.get("show_form"):
    tab1, tab2 = st.tabs(["Respondent Info", "Hazard Risk Evaluation"])
    with tab1:
        name = st.text_input("Full Name")
        final_ward = ward_display or st.text_input("Ward (if not using map)")
        today = st.date_input("Date", value=date.today())
        user_email = st.text_input("Your Email")
        confirm = st.checkbox("I confirm the information is accurate")
    with tab2:
        with st.form("hazard_form"):
            responses = []
            hazards_to_ask = selected + ([custom] if custom else [])
            for hazard in hazards_to_ask:
                st.markdown(f"### {hazard}")
                for q, opts in questions_with_descriptions.items():
                    response = st.radio(q, opts, key=f"{hazard}_{q}")
                    responses.append({"Hazard": hazard, "Question": q, "Response": response})
                st.markdown("#### Rate Capacities")
                for cq in capacity_questions:
                    response = st.radio(cq, capacity_options, key=f"{hazard}_{cq}")
                    responses.append({"Hazard": hazard, "Question": cq, "Response": response})
            if st.form_submit_button("Submit Survey"):
                if not name or not final_ward:
                    st.error("Please fill in your name and ward.")
                else:
                    csv_file, doc_file, pdf_file = save_responses(responses, name, final_ward, user_email, today)
                    st.success("Survey submitted successfully!")
                    if user_email:
                        send_email("Your KZN Hazard Survey Submission", "Thank you for completing the survey.", [user_email], [csv_file, doc_file, pdf_file])
                    send_email("New KZN Hazard Survey Submission", "A new survey has been submitted.", ADMIN_EMAILS, [csv_file, doc_file, pdf_file])
